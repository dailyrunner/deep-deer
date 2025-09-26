"""Vector store management for document and data retrieval"""
import os
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
import pickle
import fnmatch
import asyncio

from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.config import settings
from app.core.database import get_all_tables, get_table_info, engine
from app.services.embeddings import embedding_service
from sqlalchemy import text

logger = logging.getLogger(__name__)


class VectorStoreManager:
    """Manage multiple vector stores for different data sources"""

    # Supported file types
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.pdf': 'pdf',
        '.doc': 'doc',
        '.docx': 'docx',
        '.md': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.json': 'json',
        '.csv': 'csv',
        '.py': 'python',
        '.js': 'javascript',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.h': 'header',
        '.hpp': 'header',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.xml': 'xml',
        '.log': 'log',
        '.rtf': 'rtf',
        '.odt': 'odt',
        '.ppt': 'ppt',
        '.pptx': 'pptx'
    }

    def __init__(self):
        self.store_path = Path(settings.vector_store_path)
        self.store_path.mkdir(parents=True, exist_ok=True)

        self.stores: Dict[str, FAISS] = {}
        self.is_ready = False

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    async def initialize(self):
        """Initialize vector store manager"""
        try:
            # Initialize embedding service
            embedding_service.initialize()

            # Load existing stores
            await self._load_existing_stores()

            self.is_ready = True
            logger.info("Vector store manager initialized")
        except Exception as e:
            logger.error(f"Failed to initialize vector store manager: {e}")
            self.is_ready = False

    async def cleanup(self):
        """Cleanup resources"""
        self.stores.clear()
        self.is_ready = False

    async def _load_existing_stores(self):
        """Load existing vector stores from disk"""
        for store_dir in self.store_path.iterdir():
            if store_dir.is_dir():
                try:
                    store_name = store_dir.name
                    store = FAISS.load_local(
                        str(store_dir),
                        embedding_service.get_embeddings(),
                        allow_dangerous_deserialization=True
                    )
                    self.stores[store_name] = store
                    logger.info(f"Loaded vector store: {store_name}")
                except Exception as e:
                    logger.error(f"Failed to load store {store_dir.name}: {e}")

    def _save_store(self, store_name: str):
        """Save vector store to disk"""
        if store_name in self.stores:
            store_dir = self.store_path / store_name
            store_dir.mkdir(exist_ok=True)
            self.stores[store_name].save_local(str(store_dir))
            logger.info(f"Saved vector store: {store_name}")

    async def index_database_table(self, table_name: str) -> bool:
        """Index a database table into vector store"""
        try:
            # Get table data
            async with engine.connect() as conn:
                result = await conn.execute(text(f"SELECT * FROM {table_name}"))
                rows = result.fetchall()
                columns = result.keys()

            if not rows:
                logger.warning(f"No data found in table {table_name}")
                return False

            # Convert rows to documents
            documents = []
            for row in rows:
                # Create text representation of row
                content_parts = []
                metadata = {"table": table_name, "type": "database"}

                for col, value in zip(columns, row):
                    content_parts.append(f"{col}: {value}")
                    metadata[col] = str(value)

                content = "\n".join(content_parts)
                doc = Document(page_content=content, metadata=metadata)
                documents.append(doc)

            # Split documents
            split_docs = self.text_splitter.split_documents(documents)

            # Create or update vector store
            store_name = f"db_{table_name}"
            if store_name in self.stores:
                # Add to existing store
                self.stores[store_name].add_documents(split_docs)
            else:
                # Create new store
                self.stores[store_name] = FAISS.from_documents(
                    split_docs,
                    embedding_service.get_embeddings()
                )

            # Save to disk
            self._save_store(store_name)

            logger.info(f"Indexed {len(split_docs)} chunks from table {table_name}")
            return True

        except Exception as e:
            logger.error(f"Failed to index table {table_name}: {e}")
            return False

    async def index_document(self, file_path: str, store_name: str = "documents") -> bool:
        """Index a document file into vector store"""
        try:
            from pypdf import PdfReader
            import docx

            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return False

            # Extract text based on file type
            text = ""
            metadata = {
                "source": str(file_path),
                "type": "document",
                "filename": file_path.name
            }

            if file_path.suffix.lower() == '.pdf':
                reader = PdfReader(str(file_path))
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif file_path.suffix.lower() == '.docx':
                doc = docx.Document(str(file_path))
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            elif file_path.suffix.lower() in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                logger.error(f"Unsupported file type: {file_path.suffix}")
                return False

            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return False

            # Create document and split
            doc = Document(page_content=text, metadata=metadata)
            split_docs = self.text_splitter.split_documents([doc])

            # Add to vector store
            if store_name in self.stores:
                self.stores[store_name].add_documents(split_docs)
            else:
                self.stores[store_name] = FAISS.from_documents(
                    split_docs,
                    embedding_service.get_embeddings()
                )

            # Save to disk
            self._save_store(store_name)

            logger.info(f"Indexed {len(split_docs)} chunks from {file_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to index document {file_path}: {e}")
            return False

    async def index_folder(
        self,
        folder_path: str,
        store_name: str = "documents",
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
        exclude_patterns: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Index all documents in a folder

        Args:
            folder_path: Path to the folder to index
            store_name: Name of the vector store
            recursive: Process subfolders recursively
            extensions: File extensions to process (None = all supported)
            exclude_patterns: Glob patterns to exclude

        Returns:
            Statistics about indexed files
        """
        folder = Path(folder_path)
        if not folder.exists() or not folder.is_dir():
            raise ValueError(f"Invalid folder path: {folder_path}")

        # Determine extensions to process
        if extensions:
            valid_extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}'
                              for ext in extensions]
            # Filter to only supported extensions
            valid_extensions = [ext for ext in valid_extensions
                              if ext in self.SUPPORTED_EXTENSIONS]
        else:
            valid_extensions = list(self.SUPPORTED_EXTENSIONS.keys())

        # Statistics
        stats = {
            "total_files": 0,
            "indexed_files": 0,
            "failed_files": 0,
            "skipped_files": 0,
            "total_chunks": 0,
            "errors": []
        }

        # Collect all files to process
        files_to_process = []
        pattern = '**/*' if recursive else '*'

        for file_path in folder.glob(pattern):
            if not file_path.is_file():
                continue

            stats["total_files"] += 1

            # Check if file should be excluded
            if exclude_patterns:
                should_exclude = False
                for pattern in exclude_patterns:
                    if fnmatch.fnmatch(str(file_path), pattern) or \
                       fnmatch.fnmatch(file_path.name, pattern):
                        should_exclude = True
                        break
                if should_exclude:
                    stats["skipped_files"] += 1
                    logger.debug(f"Skipping excluded file: {file_path}")
                    continue

            # Check if extension is supported
            if file_path.suffix.lower() not in valid_extensions:
                stats["skipped_files"] += 1
                logger.debug(f"Skipping unsupported file type: {file_path}")
                continue

            files_to_process.append(file_path)

        # Process files in batches to avoid memory issues
        batch_size = 10
        for i in range(0, len(files_to_process), batch_size):
            batch = files_to_process[i:i + batch_size]

            # Process batch concurrently
            tasks = []
            for file_path in batch:
                tasks.append(self._index_single_file(str(file_path), store_name))

            # Wait for batch to complete
            results = await asyncio.gather(*tasks, return_exceptions=True)

            # Update statistics
            for file_path, result in zip(batch, results):
                if isinstance(result, Exception):
                    stats["failed_files"] += 1
                    stats["errors"].append({
                        "file": str(file_path),
                        "error": str(result)
                    })
                    logger.error(f"Failed to index {file_path}: {result}")
                elif result:
                    stats["indexed_files"] += 1
                    if isinstance(result, dict):
                        stats["total_chunks"] += result.get("chunks", 0)
                else:
                    stats["failed_files"] += 1

        # Save the store after all files are indexed
        if stats["indexed_files"] > 0:
            self._save_store(store_name)
            logger.info(f"Indexed folder {folder_path}: {stats['indexed_files']} files, "
                       f"{stats['total_chunks']} chunks")

        return stats

    async def _index_single_file(self, file_path: str, store_name: str) -> Dict[str, Any]:
        """Index a single file and return statistics"""
        try:
            # Load the document
            from langchain_community.document_loaders import (
                TextLoader, PyPDFLoader, Docx2txtLoader,
                CSVLoader, JSONLoader
            )
            from langchain_unstructured import UnstructuredLoader

            file_ext = Path(file_path).suffix.lower()

            # Select appropriate loader
            if file_ext == '.pdf':
                loader = PyPDFLoader(file_path)
            elif file_ext in ['.doc', '.docx']:
                loader = Docx2txtLoader(file_path)
            elif file_ext == '.csv':
                loader = CSVLoader(file_path)
            elif file_ext == '.json':
                loader = JSONLoader(file_path, jq_schema='.')
            elif file_ext in ['.txt', '.md', '.log']:
                loader = TextLoader(file_path)
            else:
                # Try generic loader for other formats
                loader = UnstructuredLoader(file_path)

            documents = loader.load()

            # Add file metadata
            for doc in documents:
                doc.metadata.update({
                    "source": file_path,
                    "file_name": Path(file_path).name,
                    "file_type": self.SUPPORTED_EXTENSIONS.get(file_ext, "unknown")
                })

            # Split documents
            split_docs = self.text_splitter.split_documents(documents)

            # Add to vector store
            if store_name in self.stores:
                self.stores[store_name].add_documents(split_docs)
            else:
                self.stores[store_name] = FAISS.from_documents(
                    split_docs,
                    embedding_service.get_embeddings()
                )

            return {"success": True, "chunks": len(split_docs)}

        except Exception as e:
            logger.error(f"Error indexing file {file_path}: {e}")
            raise e

    async def search(
        self,
        query: str,
        store_names: Optional[List[str]] = None,
        k: int = 5,
        filter_dict: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Search across vector stores"""
        if not self.is_ready:
            logger.error("Vector store manager is not ready")
            return []

        # Determine which stores to search
        if store_names:
            stores_to_search = [s for s in store_names if s in self.stores]
        else:
            stores_to_search = list(self.stores.keys())

        if not stores_to_search:
            logger.warning("No vector stores available for search")
            return []

        all_results = []

        for store_name in stores_to_search:
            try:
                store = self.stores[store_name]

                # Perform similarity search
                if filter_dict:
                    # Search with metadata filter
                    results = store.similarity_search(
                        query,
                        k=k,
                        filter=filter_dict
                    )
                else:
                    results = store.similarity_search(query, k=k)

                # Add store name to metadata
                for doc in results:
                    doc.metadata["vector_store"] = store_name

                all_results.extend(results)

            except Exception as e:
                logger.error(f"Error searching store {store_name}: {e}")

        # Sort by relevance (assuming first results are most relevant)
        # and limit to k total results
        return all_results[:k]

    async def search_with_score(
        self,
        query: str,
        store_names: Optional[List[str]] = None,
        k: int = 5,
        score_threshold: float = 0.5
    ) -> List[tuple[Document, float]]:
        """Search with similarity scores"""
        if not self.is_ready:
            logger.error("Vector store manager is not ready")
            return []

        # Determine which stores to search
        if store_names:
            stores_to_search = [s for s in store_names if s in self.stores]
        else:
            stores_to_search = list(self.stores.keys())

        all_results = []

        for store_name in stores_to_search:
            try:
                store = self.stores[store_name]
                results = store.similarity_search_with_score(query, k=k)

                # Filter by score threshold and add store name
                for doc, score in results:
                    if score >= score_threshold:
                        doc.metadata["vector_store"] = store_name
                        all_results.append((doc, score))

            except Exception as e:
                logger.error(f"Error searching store {store_name}: {e}")

        # Sort by score (higher is better) and limit to k
        all_results.sort(key=lambda x: x[1], reverse=True)
        return all_results[:k]

    def list_stores(self) -> List[str]:
        """List available vector stores"""
        return list(self.stores.keys())

    def get_store_info(self, store_name: str) -> Optional[Dict[str, Any]]:
        """Get information about a specific store"""
        if store_name not in self.stores:
            return None

        store = self.stores[store_name]
        # Note: FAISS doesn't directly expose document count
        # This is an approximation based on the index
        return {
            "name": store_name,
            "type": "FAISS",
            "embedding_model": settings.embedding_model,
            "path": str(self.store_path / store_name)
        }

    async def delete_store(self, store_name: str) -> bool:
        """Delete a vector store"""
        try:
            if store_name in self.stores:
                del self.stores[store_name]

            store_dir = self.store_path / store_name
            if store_dir.exists():
                import shutil
                shutil.rmtree(store_dir)

            logger.info(f"Deleted vector store: {store_name}")
            return True

        except Exception as e:
            logger.error(f"Failed to delete store {store_name}: {e}")
            return False


# Global instance
vector_manager = VectorStoreManager()